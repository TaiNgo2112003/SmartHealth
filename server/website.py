from flask import render_template, redirect, url_for, request, session, jsonify, send_file
from flask_babel import refresh
from time import time
from os import urandom
from server.babel import get_locale, get_languages
import pyodbc
from docx import Document
from io import BytesIO
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
from flask import flash
import pandas as pd


class Website:
    def __init__(self, bp, url_prefix) -> None:
        self.bp = bp
        self.url_prefix = url_prefix
        self.routes = {
            '/': {
                'function': lambda: redirect(url_for('._index')),
                'methods': ['GET', 'POST']
            },
            '/chat/': {
                'function': self._index,
                'methods': ['GET', 'POST']
            },
            '/chat/<conversation_id>': {
                'function': self._chat,
                'methods': ['GET', 'POST']
            },
            '/change-language': {
                'function': self.change_language,
                'methods': ['POST']
            },
            '/get-locale': {
                'function': self.get_locale,
                'methods': ['GET']
            },
            '/get-languages': {
                'function': self.get_languages,
                'methods': ['GET']
            },
            '/appointment_booking': {
                'function': self.appointment_booking,
                'methods': ['GET', 'POST']
            },
            '/diagnosis': {
                'function': self.diagnosis,
                'methods': ['GET', 'POST']
            },
            '/medical': {
                'function': self.medical_products,
                'methods': ['GET']
            },
            '/add_to_invoice': {
                'function': self.add_to_invoice,
                'methods': ['POST']
            },
            '/invoice': {
                'function': self.invoice_page,
                'methods': ['GET']
            },
            '/confirm_invoice': {
                'function': self.confirm_invoice,
                'methods': ['POST']
            }
        }
        self.conn = pyodbc.connect(
            'DRIVER={SQL Server};'
            'SERVER=TAIPC\\SQLSERVERM1;'
            'DATABASE=MedicalGPT;'  
            'UID=sa;'
            'PWD=123'
        )
        self.invoice = []

    def _chat(self, conversation_id):
        if '-' not in conversation_id:
            return redirect(url_for('._index'))

        return render_template('index.html', chat_id=conversation_id, url_prefix=self.url_prefix)

    def _index(self):
        chat_id = f'{urandom(4).hex()}-{urandom(2).hex()}-{urandom(2).hex()}-{urandom(2).hex()}-{hex(int(time() * 1000))[2:]}'
        return render_template('index.html', chat_id=chat_id, url_prefix=self.url_prefix)

    def change_language(self):
        data = request.get_json()
        session['language'] = data.get('language')
        refresh()
        return '', 204

    def get_locale(self):
        return get_locale()
    
    def get_languages(self):  
        return get_languages()

    def medical_products(self):
        search_query = request.args.get('search', '')
        cursor = self.conn.cursor()
        if search_query:
            cursor.execute("SELECT * FROM medical WHERE name LIKE ?", f'%{search_query}%')
        else:
            cursor.execute("SELECT * FROM medical")
        products = cursor.fetchall()
        return render_template('medical.html', products=products)

    def add_to_invoice(self):
        data = request.get_json()
        product_id = data['productId']
        quantity = int(data['quantity'])
        
        cursor = self.conn.cursor()
        cursor.execute("SELECT * FROM medical WHERE id = ?", product_id)
        product = cursor.fetchone()
        
        if product and quantity > 0:
            total_price = product.unit_price * quantity
            self.invoice.append({
                'id': product.id,
                'name': product.name,
                'quantity': quantity,
                'total_price': total_price
            })

        return jsonify({'status': 'success'})

    def invoice_page(self):
        return render_template('invoice.html', invoice_items=self.invoice)

    def appointment_booking(self):
        if request.method == 'GET':
            cursor = self.conn.cursor()
            cursor.execute("SELECT id, name, specialization FROM doctor")
            doctors = cursor.fetchall()
            cursor.close()
            return render_template('appointment.html', doctors=doctors)

        elif request.method == 'POST':
            try:
                patient_name = request.form.get('patient_name')
                doctor_id = request.form.get('doctor_id')
                appointment_date = request.form.get('appointment_date')
                appointment_time = request.form.get('appointment_time')
                reason_for_visit = request.form.get('reason_for_visit')
                notes = request.form.get('notes')

                if not patient_name or not doctor_id or not appointment_date or not appointment_time:
                    return jsonify({'status': 'error', 'message': 'All fields are required'}), 400

                cursor = self.conn.cursor()
                cursor.execute("""
                    SELECT COUNT(*) 
                    FROM appointment 
                    WHERE doctor_id = ? 
                    AND appointment_date = ? 
                    AND appointment_time = ?
                """, (doctor_id, appointment_date, appointment_time))
                doctor_busy = cursor.fetchone()[0]

                if doctor_busy > 0:
                    return jsonify({'status': 'error', 'message': 'Doctor is not available at this time'}), 400

                cursor.execute("""
                    INSERT INTO appointment (doctor_id, patient_name, appointment_date, appointment_time, reason_for_visit, notes)
                    VALUES (?, ?, ?, ?, ?, ?)
                """, (doctor_id, patient_name, appointment_date, appointment_time, reason_for_visit, notes))

                self.conn.commit()
                cursor.close()

                return jsonify({'status': 'success', 'message': 'Appointment booked successfully'}), 200

            except Exception as e:
                self.conn.rollback()
                return jsonify({'status': 'error', 'message': str(e)}), 500
   
    def diagnosis(self):
        if request.method == 'GET':
            return render_template('diagnosis.html')

        elif request.method == 'POST':
            data = request.get_json()
            symptoms = data.get('symptoms', [])

            # Kiểm tra số lượng triệu chứng nhập vào
            if len(symptoms) < 2 or len(symptoms) > 5:
                return jsonify({'error': 'Vui lòng nhập từ 2 đến 5 triệu chứng.'}), 400

            diagnosis_result = self.analyze_symptoms(symptoms)
            return jsonify({'predicted_disease': diagnosis_result})

    def analyze_symptoms(self, symptoms):
        try:
            df = pd.read_csv(r'C:\Users\taing\OneDrive\Máy tính\FreeGPT-4\data_pivoted.csv', index_col=0)
            print("Đã tải file CSV thành công.")
        except FileNotFoundError:
            return ["Không tìm thấy file CSV. Vui lòng kiểm tra lại đường dẫn."]
        except pd.errors.EmptyDataError:
            return ["File CSV rỗng. Vui lòng kiểm tra lại nội dung của file."]
        except Exception as e:
            return [f"Có lỗi xảy ra: {str(e)}"]

        if df.empty:
            return ["File CSV không có dữ liệu."]

        diagnosis_results = {}

        # Duyệt qua các triệu chứng đã nhập
        for symptom in symptoms:
            # Kiểm tra triệu chứng có trong DataFrame không
            if symptom in df.index:
                matched_diseases = df.columns[(df.loc[symptom].notnull())]  # Chỉ chọn các bệnh có triệu chứng

                # Cập nhật kết quả cho các bệnh đã tìm thấy
                for disease in matched_diseases:
                    diagnosis_results[disease] = diagnosis_results.get(disease, 0) + 1

        # Nếu không tìm thấy bệnh nào phù hợp với triệu chứng đã cho
        if not diagnosis_results:
            # Trả về bệnh đầu tiên trong DataFrame nếu không có bệnh nào khớp
            return [df.columns[0]]  # Trả về bệnh đầu tiên trong danh sách bệnh

        # Chỉ trả về bệnh có số lần khớp triệu chứng cao nhất
        predicted_disease = max(diagnosis_results, key=diagnosis_results.get)

        return [predicted_disease]  # Chỉ trả về bệnh phù hợp nhất

    def confirm_invoice(self):
        cursor = self.conn.cursor()
        medication_time = request.form.getlist('medication_time')
        dosage = request.form.get('dosage')
        food_time = request.form.get('food_time')
        revisit_date = request.form.get('revisit_date')
        note = request.form.get('note')
        try:
            doc = Document()
            doc.add_paragraph("Cộng hòa – Xã hội – Chủ nghĩa – Việt Nam")
            doc.add_paragraph("Độc lập – Tự do – Hạnh phúc")
            doc.add_paragraph("\nHÓA ĐƠN THUỐC\n")
            
            for item in self.invoice:
                cursor.execute("SELECT quantity FROM medical WHERE id = ?", (item['id'],))
                current_quantity = cursor.fetchone()

                if current_quantity is None:
                    return jsonify({'status': 'error', 'message': f'Item with id {item["id"]} not found'}), 404

                current_quantity = current_quantity[0]

                if current_quantity >= item['quantity']:
                    cursor.execute("UPDATE medical SET quantity = quantity - ? WHERE id = ?", (item['quantity'], item['id']))
                    doc.add_paragraph(f"Tên thuốc: {item['name']} - Số lượng: {item['quantity']}")
                else:
                    return jsonify({'status': 'error', 'message': f'Insufficient stock for item with id {item["id"]}'})

            doc.add_heading('Thông tin uống thuốc', level=1)
            doc.add_paragraph('Thời gian uống thuốc: ' + ', '.join(medication_time))
            if food_time == "before_food":
                doc.add_paragraph('Nhắc nhở uống thuốc: trước khi ăn')
            elif food_time == "after_food":
                doc.add_paragraph('Nhắc nhở uống thuốc: sau khi ăn')
            else:
                doc.add_paragraph('Không có nhắc nhở đặc biệt.')
            
            if revisit_date:
                doc.add_paragraph('Ngày hẹn tái khám: ' + revisit_date)
            if note:
                doc.add_paragraph('Ghi chú: ' + note)
            
            doc.add_paragraph("\nThời gian xuất hóa đơn: " + datetime.now().strftime("%d/%m/%Y %H:%M:%S"))
            doc_path = f'invoices/invoice_{time()}.docx'
            doc.save(doc_path)

            self.invoice.clear()  # Clear the invoice after confirmation
            return send_file(doc_path, as_attachment=True)

        except Exception as e:
            return jsonify({'status': 'error', 'message': str(e)}), 500

    def register_routes(self):
        for route, options in self.routes.items():
            self.bp.route(route, methods=options['methods'])(options['function'])
